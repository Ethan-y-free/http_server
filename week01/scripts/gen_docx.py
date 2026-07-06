"""将 ppt-design.md 转为格式化 Word 文档"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

# 颜色常量
ACCENT_BLUE = RGBColor(0x60, 0xA5, 0xFA)
ACCENT_GREEN = RGBColor(0x34, 0xD3, 0x99)
ACCENT_PINK = RGBColor(0xF4, 0x72, 0xB6)
DARK_BG = RGBColor(0x0F, 0x0F, 0x14)
CARD_BG = RGBColor(0x1A, 0x1A, 0x24)
GRAY = RGBColor(0x71, 0x71, 0x7A)
WHITE = RGBColor(0xD4, 0xD4, 0xD8)

doc = Document()

# 页边距
for section in doc.sections:
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_title(text, size=22, color=ACCENT_BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_subtitle(text, size=13, color=GRAY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_body(text, size=11, color=RGBColor(0x33, 0x33, 0x33)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_code(text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0xA5, 0xB4, 0xFC)
    run.font.name = 'Consolas'
    return p

def add_bullet(text, indent=0, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.left_indent = Cm(1 + indent * 1.5)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()  # spacing
    return table

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('─' * 60)
    run.font.color.rgb = GRAY
    run.font.size = Pt(8)

def add_highlight_box(text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    # Use a border box via shading
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F0F4FF',
        qn('w:val'): 'clear',
    })
    pPr.insert(0, shd)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


# ======================================================================
# 正文内容
# ======================================================================

add_title('高并发 HTTP 服务器 — 架构回顾 PPT 设计文档（共 11 页）')
add_subtitle('从 OS 原语到主从 Reactor · 九层架构逐层拆解')
add_subtitle('2026/07 · 4 周开发 · 压测 6234 QPS · 0 失败')
add_separator()

add_title('全局样式', size=14, color=ACCENT_GREEN)
add_bullet('比例：16:9 宽屏')
add_bullet('配色：深色背景（#0F0F14），卡片（#1A1A24），蓝色标题（#60A5FA），绿色代码（#34D399），粉色强调（#F472B6）')
add_bullet('字体：标题 28-36pt，正文 14-16pt，代码 12-13pt Consolas')
add_separator()

# ---- 第 1 页 ----
add_title('第 1 页：封面')
add_bullet('标题：高并发 HTTP 服务器 — 架构全景回顾', size=11)
add_bullet('副标题：从 OS 原语到主从 Reactor · 九层架构逐层拆解', size=11)
add_bullet('标签：C++17 | epoll | Reactor | One Loop Per Thread', size=11)
add_bullet('右下角：2026/07 · 4 周开发 · 压测 6234 QPS · 0 失败', size=11)
add_separator()

# ---- 第 2 页 ----
add_title('第 2 页：九层架构全景图')
add_subtitle('纵向分层图，从上到下：')
add_code('⑨ MultiReactorServer     — 主从调度编排（黄色背景块）')
add_code('⑧ HttpParser / TimerWheel / AsyncLogger — HTTP 业务层（黄色背景块）')
add_code('⑦ EventLoop              — One Loop Per Thread（蓝色背景块）')
add_code('⑥ ThreadPool             — 任务队列，暂未使用（蓝色背景块）')
add_code('⑤ Channel                — (fd, events, callbacks) 三元组（蓝色背景块）')
add_code('④ Buffer                 — TCP 字节流组装工厂（蓝色背景块）')
add_code('③ Epoll                  — epoll RAII 封装（绿色背景块）')
add_code('② Socket                 — socket RAII 封装（绿色背景块）')
add_code('① OS 原语               — epoll / eventfd / timerfd / readv（粉色背景块）')
add_body('层与层之间用 ↑ 箭头连接。')
add_separator()

# ---- 第 3 页 ----
add_title('第 3 页：RAII 封装层（Socket & Epoll）')
add_subtitle('移动不拷贝 · 谁创建谁释放')
add_body('布局：左右双栏对比。')
add_table(['栏位', '内容'], [
    ['左栏（问题）', '忘记 close → fd 泄漏 → "Too many open files"'],
    ['左栏（问题）', '拷贝 socket → 两个对象持同一个 fd → double close → 未定义行为'],
    ['右栏（解法）', '构造获取资源，析构自动释放'],
    ['右栏（解法）', '移动接管所有权（fd_ = other.fd_; other.fd_ = -1）'],
    ['右栏（解法）', '禁止拷贝（= delete）'],
    ['右栏（解法）', '两个关键场景：accept() 返回值移动 + shared_ptr 跨线程传递'],
])
add_body('底部要点：explicit 防隐式转换 | SO_REUSEADDR 快速重启 | SetNonBlocking 配合 epoll ET')
add_separator()

# ---- 第 4 页 ----
add_title('第 4 页：Buffer 应用层缓冲区')
add_subtitle('TCP 字节流的"组装工厂" — 非阻塞 IO 为什么必须有应用层缓冲区')
add_table(['问题', '无 Buffer', '有 Buffer'], [
    ['TCP 字节流无边界', '半条请求无法处理，下次 read 接不上', '攒够一条完整消息再交出'],
    ['非阻塞 IO 不保证一次读完', '栈上 buf 离开作用域即丢', '持久化，下次 EPOLLIN 继续 append'],
])
add_body('readv + iovec 双缓冲：')
add_code('[fd] → readv(iov, 2) → [内部 buffer_] + [栈 extrabuf 64KB]')
add_code('一次系统调用完成读取')
add_code('  - 数据 ≤ WritableBytes → 只填内部 buffer')
add_code('  - 数据 > WritableBytes → 内部填满 + 溢出到 extrabuf → 再 Append 回来')
add_body('kCheapPrepend = 8，预留头部空间，确保 prepend 不重新分配。')
add_separator()

# ---- 第 5 页 ----
add_title('第 5 页：Channel 事件分发器')
add_subtitle('(fd, events, callbacks) 三元组 — EventLoop 只管"有事件 → 调这个 Channel"')
add_body('核心设计：')
add_code('Channel ch(fd, epoll);')
add_code('ch.SetReadCallback(  [...]{ OnRead(fd);  } );')
add_code('ch.SetWriteCallback( [...]{ OnWrite(fd); } );')
add_code('ch.SetCloseCallback( [...]{ OnClose(fd); } );')
add_code('ch.SetErrorCallback( [...]{ OnError(fd); } );')
add_code('ch.EnableRead();  // events_ |= EPOLLIN | EPOLLRDHUP')
add_code('ch.EnableWrite(); // events_ |= EPOLLOUT')
add_body('事件分发优先级（HandleEvent 内部）：')
add_bullet('EPOLLHUP（无数据未读）→ closeCallback_')
add_bullet('EPOLLERR / EPOLLHUP → errorCallback_')
add_bullet('EPOLLRDHUP（对端半关闭）→ closeCallback_')
add_bullet('EPOLLIN → readCallback_')
add_bullet('EPOLLOUT → writeCallback_')
add_highlight_box('对比：没有 Channel → EventLoop 里 switch(fd) + if(events) 散落各处 → 无法维护')
add_separator()

# ---- 第 6 页 ----
add_title('第 6 页：EventLoop 事件循环核心')
add_subtitle('One Loop Per Thread + eventfd 唤醒 + swap 技巧')
add_body('核心循环：')
add_code('while (!quit_) {')
add_code('    nfds = epoll_wait(..., 64);')
add_code('    for (i = 0..nfds) channels_[fd]->HandleEvent(events);')
add_code('    DoPendingFunctors();')
add_code('}')
add_body('三个关键设计：')
add_body('① eventfd 跨线程唤醒：', size=12, color=ACCENT_BLUE)
add_bullet('没有 eventfd → quit_ = true 没人看到 → epoll_wait 永远阻塞 → 进程卡死')
add_bullet('有 eventfd → Quit() 写 8 字节 → epoll_wait 立即返回 → 安全退出')
add_bullet('作用类似"门铃"，让其他线程能唤醒本线程的 epoll_wait')
add_body('② DoPendingFunctors 的 swap 技巧：', size=12, color=ACCENT_BLUE)
add_bullet('加锁 → functors.swap(pending_functors_)（O(1) 交换 3 个指针）→ 解锁')
add_bullet('pending_functors_ 变空 → 其他线程可以安全 enqueue')
add_bullet('functors 在栈上独占 → 无竞争执行')
add_body('③ 架构价值：', size=12, color=ACCENT_BLUE)
add_bullet('无锁（fd 不跨线程共享）+ Cache 友好 + 故障隔离')
add_separator()

# ---- 第 7 页 ----
add_title('第 7 页：ThreadPool 线程池')
add_subtitle('当前未使用，但已就绪 — 为什么没用到 + 什么时候用')
add_body('为什么没用上？')
add_highlight_box('主从 Reactor 天然"一个连接一个线程负责到底"：read → HTTP 解析 → 生成响应 → write 全在 EventLoop 线程里完成，没有需要卸掉的 CPU 密集任务。')
add_body('将来什么时候用？')
add_table(['场景', '原因'], [
    ['gzip 压缩响应体', 'CPU 密集，会阻塞 EventLoop'],
    ['大文件磁盘 IO', '可能长时间阻塞'],
    ['CGI / FastCGI', '外部进程调用，等待时间长'],
])
add_body('设计亮点：')
add_bullet('RAII 析构安全关闭（stop_ + notify_all + join）')
add_bullet('Submit() 支持 std::future 获取返回值')
add_bullet('关闭时拒绝新任务（stop_ 检查 → 抛异常）')
add_separator()

# ---- 第 8 页 ----
add_title('第 8 页：HTTP 解析器 — 状态机跨 TCP 分片')
add_subtitle('TCP 字节流 + 状态机 + Buffer = 消息边界感知')
add_highlight_box('核心公式：TCP 字节流（无边界）+ 状态机（记住当前位置）+ Buffer（攒数据）= 感知消息边界')
add_body('三个状态：')
add_code('PARSE_REQUEST_LINE → 找 \\r\\n → 解析 "GET /index.html HTTP/1.1"')
add_code('PARSE_HEADERS      → 逐行解析直到空行 \\r\\n')
add_code('PARSE_BODY         → 累积 Content-Length 字节')
add_code('PARSE_DONE         → 完整的 HttpRequest 对象')
add_body('NEED_MORE 是精髓：')
add_highlight_box('数据不够时不报错，保留 Buffer 数据，保持当前状态 → 等下次 EPOLLIN 继续。')
add_table(['对比', '行为', '结果'], [
    ['不用状态机', '第一次 read "GET /inde" → 报错/丢弃', '第二次 read 接不上 → 请求丢失'],
    ['用状态机', '状态停在第 1 步，Buffer 保留', '下次 read 继续找 \\r\\n → 完整解析'],
])
add_separator()

# ---- 第 9 页 ----
add_title('第 9 页：TimerWheel 时间轮')
add_subtitle('O(1) 踢空闲连接 · 60 槽 × 1 秒 = 60 秒超时窗口 · 每个 SubReactor 独立')
add_body('核心问题：恶意客户端 TCP 连上了但不发 HTTP 请求 → 占着 fd / 内存 / 槽位 → 必须踢掉。')
add_table(['', '遍历所有连接', '时间轮'], [
    ['复杂度', 'O(N)', 'O(1)'],
    ['1 万连接', '1 万次检查/秒', '1 个槽位/秒'],
    ['10 万连接', '10 万次检查/秒', '1 个槽位/秒'],
])
add_body('活跃 vs 超时的区分：')
add_bullet('活跃连接 → 每次 HTTP 请求后 AddOrRefresh(fd, 60000) → 挪到远处 → 永不超时')
add_bullet('超时连接 → 一直停在原地 → Tick() 收割 → close(fd) + 释放资源')
add_body('集成方式：')
add_bullet('timerfd_create → 1 秒周期 → EPOLLIN 触发')
add_bullet('每个 SubReactor 独立 timerfd + TimerWheel → 无锁')
add_bullet('Flush() 复用同一个 timerfd 回调 → 不额外创建定时器')
add_separator()

# ---- 第 10 页 ----
add_title('第 10 页：AsyncLogger 异步日志系统')
add_subtitle('两层无锁设计 + 双缓冲流水线 · 4 个文件 4 个职责')
add_table(['文件', '职责'], [
    ['log_stream.h', 'LogStream — 4KB 格式化缓冲区 + operator<<'],
    ['log_buffer.h', 'LogBuffer — 双缓冲（current_ + next_）+ Flush() + swap'],
    ['async_log_writer.h', 'AsyncLogWriter — 后台线程 + 条件变量 + fwrite 写盘'],
    ['logger.h', 'Logger — RAII + 时间戳前缀 + LOG_INFO/ERROR/FATAL 宏'],
])
add_body('两层无锁设计：')
add_bullet('① 每线程独享 LogBuffer（thread_local 绑定）→ 线程间不共享 → 无锁')
add_bullet('② 双缓冲 swap 瞬间切换 current_ ↔ next_ → 前后端物理内存互不干扰')
add_body('双缓冲流水线：')
add_code('前端线程（SubReactor）          后端线程（AsyncLogWriter）')
add_code('  LOG_INFO << "xxx"               等待条件变量')
add_code('  → 写入 current_')
add_code('  → Flush() 时 swap → notify  →  收到完整 buffer')
add_code('  → 前端继续写 current_（空）      → fwrite(next_) 写盘')
add_body('Flush 精妙处：复用 timerfd 每秒回调，不额外创建定时器。连接关闭时自动触发最后一次 Flush。')
add_separator()

# ---- 第 11 页 ----
add_title('第 11 页：总结')
add_subtitle('九层架构一览 + 五大设计原则')
add_table(
    ['层', '模块', '核心职责', '一句话'],
    [
        ['⑨', 'MultiReactorServer', '主从调度', 'Accept 轮询 → SubReactor 全程处理'],
        ['⑧', 'HttpParser', 'HTTP 状态机', '跨 TCP 分片攒完整请求'],
        ['⑧', 'TimerWheel', '空闲连接回收', 'O(1) Tick · 活跃刷新远处 · 超时原地收割'],
        ['⑧', 'AsyncLogger', '双缓冲写盘', '每线程独享 buffer + 前后端流水线'],
        ['⑦', 'EventLoop', '事件循环', 'One Loop Per Thread + eventfd 唤醒'],
        ['⑥', 'ThreadPool', '任务队列', '暂未使用 · 预留 CPU 密集场景'],
        ['⑤', 'Channel', '事件分发', '(fd, events, callbacks) 三元组'],
        ['④', 'Buffer', '应用层缓冲', 'TCP 字节流 → 完整消息的组装工厂'],
        ['③②', 'Epoll / Socket', 'RAII 封装', '移动不拷贝 · 谁创建谁释放'],
        ['①', 'OS 原语', '内核', 'epoll / eventfd / timerfd / readv'],
    ]
)
add_body('')
add_highlight_box('五大设计原则：RAII 管理资源 | 移动不拷贝 | One Loop Per Thread | 无锁设计 | swap 掏空技巧')
add_body('')
add_subtitle('面试应对：讲完这一页留 1-2 分钟给面试官提问。准备好回答："你最满意哪个模块的设计？"→ 答 EventLoop 的 eventfd + swap 组合。')

# 保存
out = r'C:\Users\Lenovo\projects\http_server\week01\docs\ppt-design.docx'
doc.save(out)
print(f'Done → {out}')
